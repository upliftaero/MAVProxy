#!/usr/bin/env python
"""
  Data analysis module for analyzing data produced by MAVProxy testpilot module.
  This module generates an MS Word format flight test report.
  The testpilot module will automatically generate this report, or this module
  can be run directly from the command line by specifying a .csv file as a command line argument.

  Authors:   Mark Jacobsen
             Kevin Wells
"""

import sys
import os
import datetime
import argparse

from pandas.io.parsers import read_csv

import matplotlib.pyplot as plt
import matplotlib.image as mpimg
#from matplotlib.backends.backend_pdf import PdfPages        # Obsolete unless we move back to PDF files

from docx import Document
from docx.shared import Inches

import statsmodels.api as sm
import mavproxy_testpilot as tp            # Currently only for configuration constants

lowess = sm.nonparametric.lowess

class TPAnalyze:

    def __init__(self, directory, configuration):
        self.directory = directory
        self.configuration = configuration
        return

    def build_power_report(self, csvfile):
        cells = 4
        amps = 10

        _data = read_csv(self.directory + os.sep + csvfile)
        yvalues = _data['SYS_STATUS.current_battery']/100.0*_data['SYS_STATUS.voltage_battery']/1000.0
        xvalues = _data['VFR_HUD.airspeed']
        yline = lowess(yvalues,xvalues,return_sorted=True)

        endurance_high = amps/(yline[:,1]/(cells*4.2))
        endurance_med = amps/(yline[:,1]/(cells*3.7))
        endurance_low = amps/(yline[:,1]/(cells*3.0))

        range_high = endurance_high * 3.6 * yline[:,0]
        range_med = endurance_med * 3.6 * yline[:,0]
        range_low = endurance_low * 3.6 * yline[:,0]

        filebase = datetime.date.today().strftime("%Y-%m-%d_")              # TODO: Would like to pull the date from the log file directly
        ac_string = self.configuration[tp.KEY_AIRCRAFT] if tp.KEY_AIRCRAFT in self.configuration else "default-aircraft"
        filebase += ac_string.replace(' ', '-') + "_"
        ending = ".csv"
        if csvfile.endswith(ending):
            filebase += csvfile[:-len(ending)]
        else:
            filebase += csvfile

        report = TPReport(self.directory, filebase, self.configuration)

        #pp = PdfPages(self.directory + os.sep + filebase+".pdf") # Currently not used

        # Plot 1: Power
        plt.plot(xvalues,yvalues,'.',color='0.9')
        plt.plot(yline[:,0],yline[:,1],'r',linewidth=3.0)
        title_str = "Power Required vs Airspeed"
        plt.title(title_str)
        plt.xlabel("Airspeed (m/s)")
        plt.ylabel("Watts Required (W)")
        #pp.savefig(plt)
        self.save_chart(plt, filebase+"_power.png", report)

        # Plot 2: Current vs Airspeed
        #plt.plot(xvalues,yvalues/(cells*4.2),'.',color='0.9')
        plt.plot(yline[:,0],yline[:,1]/(cells*4.2),'g',linewidth=2.0)
        plt.plot(yline[:,0],yline[:,1]/(cells*3.7),'b',linewidth=2.0)
        plt.plot(yline[:,0],yline[:,1]/(cells*3.0),'r',linewidth=2.0)
        title_str = "Current Required vs Airspeed"
        plt.title(title_str)
        plt.xlabel("Airspeed (m/s)")
        plt.ylabel("Current Required (A)")
        plt.legend(['4.2V','3.7V','3.0V'],loc=2)
        self.save_chart(plt, filebase+"_i_vs_as.png", report)

        # Plot 3: Endurance vs Airspeed
        #plt.plot(xvalues,yvalues/(cells*4.2),'.',color='0.9')
        plt.plot(yline[:,0],endurance_high,'g',linewidth=2.0)
        plt.plot(yline[:,0],endurance_med,'b',linewidth=2.0)
        plt.plot(yline[:,0],endurance_low,'r',linewidth=2.0)
        title_str = "Endurance vs Airspeed"
        plt.title(title_str)
        plt.xlabel("Airspeed (m/s)")
        plt.ylabel("Endurance (hr)")
        plt.legend(['4.2V','3.7V','3.0V'],loc=1)
        self.save_chart(plt, filebase+"_endur_vs_as.png", report)

        # Plot 4: Range vs Airspeed
        #plt.plot(xvalues,yvalues/(cells*4.2),'.',color='0.9')
        plt.plot(yline[:,0],range_high,'g',linewidth=2.0)
        plt.plot(yline[:,0],range_med,'b',linewidth=2.0)
        plt.plot(yline[:,0],range_low,'r',linewidth=2.0)
        title_str = "Range vs Airspeed"
        plt.title(title_str)
        plt.xlabel("Airspeed (m/s)")
        plt.ylabel("Range (km)")
        plt.legend(['4.2V','3.7V','3.0V'],loc=1)
        self.save_chart(plt, filebase+"_range_vs_as.png", report)

        report.save()
        #pp.close()      # Not necessary if using "with."  See: http://matplotlib.org/examples/pylab_examples/multipage_pdf.html

        for i in range(8,35):
            print(str(i) + ", " + str(self.lowess_predict(yline,i)))

    def save_chart(self, plt, filename, report):
        fullfilename = self.directory + os.sep + filename
        plt.savefig(fullfilename, format='png')
        report.add_chart(fullfilename)
        #plt.savefig(pp,format='pdf')
        plt.close()

    def lowess_predict(self, fitted,value):
        i = 0
        if fitted[i,0] > value:
            return -1

        while value > fitted[i,0]:
            i = i + 1
            if i >= fitted.shape[0]:
                return -1
        return fitted[i,1]

class TPReport:

    def __init__(self, directory, filebase, configuration):
        self.directory = directory
        self.filebase = filebase
        self.configuration = configuration
        self.document = Document()
        logofile = self.directory + os.sep + "uplift_logo.png"
        if os.path.isfile(logofile):
            self.document.add_picture(logofile, width=Inches(6))
        self.document.add_heading('Flight Test Report', 0)
        self.document.add_paragraph(datetime.date.today().strftime("%B %d, %Y"))        # TODO: Would like to pull the date from the log file directly
        for key in configuration:
            p = self.document.add_paragraph(key + ":\t" + configuration[key])

    def add_chart(self, chart_filename):
        self.document.add_picture(chart_filename, width=Inches(6))

    def save(self):
        self.document.save(self.directory + os.sep + self.filebase + ".docx")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Analyze MAVlink log files and testpilot .csv files.')
    parser.add_argument('csvfile', metavar="csv-file", help='testpilot .csv file to analyze')
    args = parser.parse_args()
    directory = os.getcwd()
    filename = args.csvfile
    if not os.path.isfile(filename):
        print("File '" + filename + "' does not exist")
        quit()
    analyzer = TPAnalyze(directory, {})
    analyzer.build_power_report(filename)
